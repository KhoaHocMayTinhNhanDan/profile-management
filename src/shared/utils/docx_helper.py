import os
import re
from docx import Document as DocxDocument


def fill_document_content_controls(
    template_doc_path: str, data: dict, output_path: str
) -> None:
    """Fills Plain Text Content Controls in the Word template matching keys in data.

    Also falls back to replacing traditional placeholder tags (e.g. {{key}}) in text runs.
    """
    if not os.path.exists(template_doc_path):
        raise FileNotFoundError(f"Template docx file not found at: {template_doc_path}")

    doc = DocxDocument(template_doc_path)

    # 1. Helper function to process all elements matching w:sdt in an XML tree root
    def process_sdt_elements(root_elem):
        sdts = root_elem.xpath("//*[local-name()='sdt']")
        for sdt in sdts:
            tag_elems = sdt.xpath(".//*[local-name()='tag']")
            if tag_elems:
                tag_val = tag_elems[0].get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                )
                if tag_val and tag_val in data:
                    # Remove showingPlc flag to let MS Word render normal run formatting
                    sdtPr_elems = sdt.xpath(".//*[local-name()='sdtPr']")
                    if sdtPr_elems:
                        sdtPr = sdtPr_elems[0]
                        plc_elems = sdtPr.xpath(".//*[local-name()='showingPlc']")
                        for plc in plc_elems:
                            sdtPr.remove(plc)

                    text_elems = sdt.xpath(
                        ".//*[local-name()='sdtContent']//*[local-name()='t']"
                    )
                    if text_elems:
                        # Set value on the first text block, clear the rest to keep styling
                        text_elems[0].text = str(data[tag_val])
                        for t in text_elems[1:]:
                            t.text = ""

    # Process document body
    process_sdt_elements(doc.element)

    # Process all headers and footers to ensure completeness
    for section in doc.sections:
        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            if header is not None:
                h_elem = getattr(header, "_element", None)
                if h_elem is not None:
                    process_sdt_elements(h_elem)
        for footer in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            if footer is not None:
                f_elem = getattr(footer, "_element", None)
                if f_elem is not None:
                    process_sdt_elements(f_elem)

    # 2. Fallback traditional placeholder replacement {{key}} (tolerating optional spacing)
    def replace_placeholders(text, placeholder_data):
        for k, v in placeholder_data.items():
            # Match {{key}} with optional spacing e.g. {{ key }}
            pattern = r"\{\{\s*" + re.escape(k) + r"\s*\}\}"
            text = re.sub(pattern, str(v), text)
        return text

    # Replace in body paragraphs
    for p in doc.paragraphs:
        for run in p.runs:
            if run.text:
                run.text = replace_placeholders(run.text, data)

    # Replace in body tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        if run.text:
                            run.text = replace_placeholders(run.text, data)

    # Save to final output path
    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)


def extract_document_content_controls(
    doc_path: str, old_data: dict | None = None
) -> dict:
    """Extracts tag and value pairs from all Plain Text Content Controls in a Word document.

    If a tag is present multiple times, it resolves duplicates by preferring values that
    differ from the old database value (the edited fields).
    """
    from src.shared.logger.app_logger import get_logger

    logger = get_logger(__name__)

    logger.info(f"--- START EXTRACT CONTENT CONTROLS for {doc_path} ---")
    logger.info(f"Old database dynamic_data reference: {old_data}")

    if not os.path.exists(doc_path):
        logger.warning(f"File not found: {doc_path}")
        return {}

    doc = DocxDocument(doc_path)
    tag_to_values = {}

    def extract_sdt_elements(root_elem):
        sdts = root_elem.xpath("//*[local-name()='sdt']")
        for sdt in sdts:
            tag_elems = sdt.xpath(".//*[local-name()='tag']")
            if tag_elems:
                tag_val = tag_elems[0].get(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val"
                )
                if tag_val:
                    text_elems = sdt.xpath(
                        ".//*[local-name()='sdtContent']//*[local-name()='t']"
                    )
                    text_val = "".join([t.text for t in text_elems if t.text]).strip()

                    # Filter out default placeholder values
                    is_placeholder = (
                        not text_val
                        or text_val == "Click here to enter text."
                        or text_val == f"[{tag_val}]"
                    )

                    logger.info(
                        f"Found content control tag: '{tag_val}', raw value: '{text_val}', is_placeholder: {is_placeholder}"
                    )

                    if not is_placeholder:
                        if tag_val not in tag_to_values:
                            tag_to_values[tag_val] = []
                        if text_val not in tag_to_values[tag_val]:
                            tag_to_values[tag_val].append(text_val)

    # Scan body, headers, and footers
    extract_sdt_elements(doc.element)

    for section in doc.sections:
        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            if header is not None:
                h_elem = getattr(header, "_element", None)
                if h_elem is not None:
                    extract_sdt_elements(h_elem)
        for footer in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            if footer is not None:
                f_elem = getattr(footer, "_element", None)
                if f_elem is not None:
                    extract_sdt_elements(f_elem)

    logger.info(
        f"All extracted non-placeholder values before deduplication: {tag_to_values}"
    )

    # Resolve duplicate values surgically
    extracted = {}
    for tag, vals in tag_to_values.items():
        if not vals:
            continue
        if len(vals) == 1:
            extracted[tag] = vals[0]
            logger.info(f"Tag '{tag}': Single value found: '{vals[0]}'")
        else:
            if old_data and tag in old_data:
                old_val = str(old_data[tag]).strip()
                # Find the value that has changed from the database record
                new_vals = [v for v in vals if v != old_val]
                logger.info(
                    f"Tag '{tag}': Duplicate values: {vals}. Old DB value: '{old_val}'. Different values: {new_vals}"
                )
                if new_vals:
                    extracted[tag] = new_vals[0]
                    logger.info(f"Tag '{tag}': Selected edited value: '{new_vals[0]}'")
                else:
                    extracted[tag] = old_val
                    logger.info(
                        f"Tag '{tag}': No changed values, fallback to old DB value: '{old_val}'"
                    )
            else:
                extracted[tag] = vals[0]
                logger.info(
                    f"Tag '{tag}': No old data reference, fallback to first value: '{vals[0]}'"
                )

    logger.info(
        f"--- END EXTRACT CONTENT CONTROLS. Final resolved dict: {extracted} ---"
    )
    return extracted


def upgrade_docx_placeholders(doc_path: str, fields: list, output_path: str) -> None:
    """Finds plain text placeholders like {{key}} or {{ key }} in the Word file,

    and upgrades them programmatically to Plain Text Content Controls tagged with 'key'.
    """
    if not os.path.exists(doc_path):
        return

    doc = DocxDocument(doc_path)
    nsmap = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
    w_ns = nsmap["w"]

    def process_paragraph(p):
        # 1. Clean split runs first by merging if needed
        for field in fields:
            pattern = re.compile(r"\{\{\s*" + re.escape(field) + r"\s*\}\}")
            if pattern.search(p.text):
                has_single_run = False
                for run in p.runs:
                    if pattern.search(run.text):
                        has_single_run = True
                        break
                if not has_single_run:
                    if len(p.runs) > 1:
                        full_text = "".join([r.text for r in p.runs])
                        p.runs[0].text = full_text
                        for r in p.runs[1:]:
                            r.text = ""

        # 2. Find and replace with Content Controls
        runs = p.runs
        if not runs:
            return

        for field in fields:
            pattern = re.compile(r"\{\{\s*" + re.escape(field) + r"\s*\}\}")
            i = 0
            while i < len(p.runs):
                run = p.runs[i]
                text = run.text
                if not text:
                    i += 1
                    continue

                match = pattern.search(text)
                if match:
                    start, end = match.span()
                    prefix = text[:start]
                    suffix = text[end:]

                    # Update current run to hold the prefix
                    run.text = prefix

                    # Construct sdt XML node using etree
                    from lxml import etree  # type: ignore

                    sdt = etree.Element(f"{{{w_ns}}}sdt", nsmap=nsmap)
                    sdtPr = etree.SubElement(sdt, f"{{{w_ns}}}sdtPr", nsmap=nsmap)
                    tag = etree.SubElement(sdtPr, f"{{{w_ns}}}tag", nsmap=nsmap)
                    tag.set(f"{{{w_ns}}}val", field)
                    alias = etree.SubElement(sdtPr, f"{{{w_ns}}}alias", nsmap=nsmap)
                    alias.set(f"{{{w_ns}}}val", field)

                    sdtContent = etree.SubElement(
                        sdt, f"{{{w_ns}}}sdtContent", nsmap=nsmap
                    )
                    r_node = etree.SubElement(sdtContent, f"{{{w_ns}}}r", nsmap=nsmap)
                    t_node = etree.SubElement(r_node, f"{{{w_ns}}}t", nsmap=nsmap)
                    t_node.text = f"[{field}]"

                    # Insert sdt after current run in paragraph XML element
                    run_el = getattr(run, "_r", None)
                    if run_el is not None:
                        parent = run_el.getparent()
                        if parent is not None:
                            idx = parent.index(run_el)
                            parent.insert(idx + 1, sdt)

                            if suffix:
                                new_r_el = etree.Element(f"{{{w_ns}}}r", nsmap=nsmap)
                                new_t_el = etree.SubElement(
                                    new_r_el, f"{{{w_ns}}}t", nsmap=nsmap
                                )
                                new_t_el.text = suffix
                                parent.insert(idx + 2, new_r_el)

                    # Reset and loop again since structure changed
                    i = 0
                    continue
                i += 1

    # Process paragraphs, tables, headers, footers
    for p in doc.paragraphs:
        process_paragraph(p)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_paragraph(p)

    for section in doc.sections:
        for header in (
            section.header,
            section.first_page_header,
            section.even_page_header,
        ):
            if header is not None:
                for p in header.paragraphs:
                    process_paragraph(p)
                for t in header.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                process_paragraph(p)
        for footer in (
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            if footer is not None:
                for p in footer.paragraphs:
                    process_paragraph(p)
                for t in footer.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                process_paragraph(p)

    out_dir = os.path.dirname(output_path)
    if out_dir:
        os.makedirs(out_dir, exist_ok=True)
    doc.save(output_path)
