from src.layer_03_interface_adapters.gateways.outbound.i_generate_document_from_template_data_source import IGenerateDocumentFromTemplateDataSource
from src.layer_04_infrastructure.databases.sqlite.sqlite_document_store import SqliteDocumentStore
from typing import Optional
from docx import Document as DocxDocument
import os

class SqliteGenerateDocumentFromTemplateDataSource(IGenerateDocumentFromTemplateDataSource):
    def __init__(self, store: Optional[SqliteDocumentStore] = None):
        self._store = store if store is not None else SqliteDocumentStore()

    async def get_profile(self, profile_id: str) -> Optional[dict]:
        return self._store.get_document("profiles", profile_id)

    async def save(self, data: dict) -> None:
        profile_id = data.get("profile_id", "")
        self._store.set_document("profiles", profile_id, data)

    async def generate_docx(self, template_doc_path: str, data: dict, output_path: str) -> str:
        if not os.path.exists(template_doc_path):
            raise FileNotFoundError(f"Template docx file not found at: {template_doc_path}")

        # Load Docx
        doc = DocxDocument(template_doc_path)

        def replace_placeholders(text, placeholder_data):
            for k, v in placeholder_data.items():
                target = f"{{{{{k}}}}}"
                if target in text:
                    text = text.replace(target, str(v))
            return text

        # Replace in paragraphs
        for p in doc.paragraphs:
            for run in p.runs:
                if run.text:
                    run.text = replace_placeholders(run.text, data)

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        for run in p.runs:
                            if run.text:
                                run.text = replace_placeholders(run.text, data)

        # Save to output path
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        doc.save(output_path)

        # Return file:// absolute path representing the URL
        return "file:///" + os.path.abspath(output_path).replace("\\", "/")
